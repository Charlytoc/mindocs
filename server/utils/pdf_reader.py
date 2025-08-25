# utils/document_reader.py
import subprocess
import re
from abc import ABC, abstractmethod
import os
import hashlib
import fitz  # PyMuPDF
from fitz import Document as FPDFDocument
from server.utils.printer import Printer
import base64
from PIL import Image
from server.ai.ai_interface import AIInterface
from docx import Document
from docxtpl import DocxTemplate

PAGE_CONNECTOR = "\n---PAGE---\n"


# =========================
# Base strategy
# =========================
printer = Printer("PDF_READER")


class DocumentStrategy(ABC):
    document_hash: str | None = None

    @abstractmethod
    def read(self, path: str) -> str:
        pass

    def hash_text(self, text: str) -> str:
        return hashlib.sha256(text.encode("utf-8")).hexdigest()

    def split_pages(self, text: str) -> list[str]:
        return text.split(PAGE_CONNECTOR)


# =========================
# Estrategias específicas
# =========================


def get_base64_image(page: fitz.Page) -> str:
    pix = page.get_pixmap()
    img = Image.open(io.BytesIO(pix.tobytes()))
    buffered = io.BytesIO()
    img.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
    return img_str


class PyMuPDFWithOCRStrategy(DocumentStrategy):
    SAMPLE_PAGES = 3
    MAX_OCR_PAGES_PER_READ = 7

    def __init__(self):
        self.ai = AIInterface(
            provider=os.getenv("PROVIDER", "ollama"),
            api_key=os.getenv("PROVIDER_API_KEY", "asdasd"),
            base_url=os.getenv("PROVIDER_BASE_URL", None),
        )

    def read(self, path: str) -> str:
        pages = []
        with fitz.open(path, filetype="pdf") as pdf:
            what_to_do = self.select_strategy(pdf)

            if what_to_do == "OCR":
                # Procesar en lotes de MAX_OCR_PAGES_PER_READ
                total_pages = pdf.page_count
                for i in range(0, total_pages, self.MAX_OCR_PAGES_PER_READ):
                    batch = []
                    page_indices = range(
                        i, min(i + self.MAX_OCR_PAGES_PER_READ, total_pages)
                    )
                    for page_idx in page_indices:
                        page = pdf[page_idx]
                        img_str = get_base64_image(page)
                        batch.append(
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{img_str}"
                                },
                            }
                        )
                    # Un solo mensaje con varias imágenes
                    res = self.ai.chat(
                        model=os.getenv("MODEL", "gemma3"),
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": (
                                            "Extrae el texto de CADA imagen adjunta. Si hay elementos gráficos asociados, explícalo."
                                            "Devuelve un bloque por página, en el mismo orden. "
                                            "Si una imagen no tiene texto, explícalo."
                                        ),
                                    },
                                    *batch,
                                ],
                            }
                        ],
                    )
                    # Suponiendo que la respuesta contiene bloques de texto separados por marcador
                    batch_texts = res.choices[0].message.content.split(PAGE_CONNECTOR)
                    pages.extend(batch_texts)
            else:
                # Estrategia normal, página por página
                for page in pdf:
                    text = page.get_text()
                    pages.append(text)
        return PAGE_CONNECTOR.join(pages)

    def select_strategy(self, sample: FPDFDocument):
        printer.yellow(f"Number of pages for PDF: {sample.page_count}")
        sample_count = min(sample.page_count, self.SAMPLE_PAGES)
        printer.yellow(f"Sample count: {sample_count}")

        sample_results = []

        for page in sample.pages(0, sample_count):
            text_result = page.get_text()
            img_str = get_base64_image(page)
            res = self.ai.chat(
                model=os.getenv("MODEL", "gemma3"),
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extrae únicamente el texto presente en la imagen, si no hay texto, devuelve un texto vacío que diga: 'EMPTY'.",
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{img_str}"
                                },
                            },
                        ],
                    },
                ],
            )
            ocr_result = res.choices[0].message.content

            if len(text_result) > len(ocr_result):
                sample_results.append("TEXT")
            else:
                sample_results.append("OCR")

        printer.yellow(f"Sample results: {sample_results}")

        if sample_results.count("OCR") > sample_results.count("TEXT"):
            return "OCR"
        else:
            return "TEXT"


class DocxStrategy(DocumentStrategy):
    def read(self, path: str) -> str:
        doc = Document(path)
        parts = []

        # Encabezado (header) principal
        header = doc.sections[0].header
        if header and header.paragraphs:
            for para in header.paragraphs:
                if para.text.strip():
                    parts.append(f"{para.text.strip()}")

        # Cuerpo principal (párrafos y tablas, en orden)
        def iter_block_items(parent):
            from docx.table import _Cell, Table
            from docx.text.paragraph import Paragraph
            for child in parent.element.body.iterchildren():
                if child.tag.endswith('tbl'):
                    yield Table(child, parent)
                elif child.tag.endswith('p'):
                    yield Paragraph(child, parent)

        for block in iter_block_items(doc):
            from docx.table import Table
            from docx.text.paragraph import Paragraph

            if isinstance(block, Paragraph):
                txt = block.text.strip()
                if txt:
                    # Encabezados (si aplica)
                    style = getattr(block.style, 'name', '').lower()
                    if style.startswith('heading'):
                        level = ''.join(filter(str.isdigit, style)) or "1"
                        parts.append(f"{'#' * int(level)} {txt}")
                    else:
                        parts.append(txt)
            elif isinstance(block, Table):
                md_table = []
                for i, row in enumerate(block.rows):
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    md_table.append(" | ".join(cells))
                    if i == 0:
                        md_table.append(" | ".join(['---'] * len(cells)))
                parts.append('\n'.join(md_table))

        # Pies de página (footer)
        footer = doc.sections[0].footer
        if footer and footer.paragraphs:
            for para in footer.paragraphs:
                if para.text.strip():
                    parts.append(f"> {para.text.strip()}")

        return "\n\n".join(parts)

class MarkdownStrategy(DocumentStrategy):
    def read(self, path: str) -> str:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()


class DocPandocStrategy(DocumentStrategy):
    def read(self, path: str) -> str:
        result = subprocess.run(
            ["pandoc", path, "-t", "plain"],
            capture_output=True,
            text=True,
            check=True,
        )
        return result.stdout


# =========================
# Lector de documentos
# =========================


class DocumentReader:
    text: str | None = None

    def __init__(self):
        self.strategy: DocumentStrategy | None = None

    def _get_strategy(self, path: str) -> DocumentStrategy:
        ext = os.path.splitext(path)[1].lower()

        if ext == ".pdf":
            return PyMuPDFWithOCRStrategy()
        elif ext == ".docx":
            return DocxStrategy()
        elif ext == ".md":
            return MarkdownStrategy()
        elif ext == ".doc":
            return DocPandocStrategy()
        else:
            raise ValueError(f"Tipo de archivo '{ext}' no soportado")

    def read(self, path: str) -> str:
        if not os.path.isfile(path):
            raise FileNotFoundError(f"Archivo no encontrado: {path}")

        self.strategy = self._get_strategy(path)
        self.text = self.strategy.read(path)

        return self.text

    def split_pages(self, text: str) -> list[str]:
        if self.strategy is None:
            raise ValueError("No se ha leído ningún documento todavía")
        return self.strategy.split_pages(text)

    def get_hash(self) -> str:
        if self.text is None:
            raise ValueError("No se ha leído ningún documento todavía")
        return self.strategy.hash_text(self.text)


def find_placeholders(text: str) -> list[str]:
    return re.findall(r"\{\{\s*([^\{\}\s]+)\s*\}\}", text)




def generate_docx_from_template(template_path: str, variables: dict, output_path: str = None) -> str:
    """
    Generates a new .docx file from a template, replacing placeholders with variable values.

    Args:
        template_path (str): Path to the .docx template file.
        variables (dict): Dictionary with variable_name: variable_value.
        output_path (str, optional): Path to save the new file. If not provided, a new file will be created in the same directory.

    Returns:
        str: Path to the generated .docx file.
    """
    doc = DocxTemplate(template_path)
    doc.render(variables)

    if not output_path:
        base, ext = os.path.splitext(template_path)
        output_path = f"{base}_generated{ext}"

    doc.save(output_path)
    return output_path


def docx_to_html(path: str) -> str:
    """
    Converts a .docx file to HTML using pandoc, returns the HTML as a string.
    """
    result = subprocess.run(
        ["pandoc", path, "-t", "html"],
        capture_output=True,
        text=True,
        check=True,
    )
    return result.stdout
