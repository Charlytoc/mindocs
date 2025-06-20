from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título centrado
title = doc.add_heading('Título del Documento', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Texto normal
doc.add_paragraph('Este es el cuerpo principal del documento.')

# Texto en la esquina superior derecha (usando alineación)
p = doc.add_paragraph()
run = p.add_run('Texto en la esquina superior derecha')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Espacio para la firma (al final, alineado a la derecha)
doc.add_paragraph('\n\n\n')  # Espacio vertical
firma = doc.add_paragraph('__________________________')
firma.alignment = WD_ALIGN_PARAGRAPH.RIGHT
firma2 = doc.add_paragraph('Firma')
firma2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Guardar el documento
doc.save('ejemplo.docx')
